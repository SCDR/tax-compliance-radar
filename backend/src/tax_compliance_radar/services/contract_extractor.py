"""合同文件字段抽取服务。

流程：
1. `extract_text_from_file` —— 根据文件扩展名把 PDF / DOCX / 文本 转换为纯文本。
2. `extract_audit_fields` —— 拿到 `countries.yaml` 里定义的 business_fields，构造 prompt 让 LLM
   输出严格 JSON，随后做类型/枚举兜底校验，返回可以直接喂给前端 `auditForm.setFieldsValue`
   的数据结构。

设计目标：
- 仅识别 `CountryRegistry` 中已定义的国家和字段，避免污染表单。
- multiselect / select 的取值必须落在 options 白名单里，未命中则丢弃。
- 抽取失败 / LLM 输出无法解析时，抛出带明确 message 的 ValueError，由路由层转成 4xx。
"""

from __future__ import annotations

import io
import json
import logging
import re
from typing import Any

from tax_compliance_radar.registry.registry import CountryRegistry
from tax_compliance_radar.services.llm_service import _achat_with_fallback, _extract_json

logger = logging.getLogger(__name__)

MAX_TEXT_CHARS = 20000  # LLM 上下文预算保护，20k 字符对普通合同足够


# ---------- 文本抽取 ----------

def _extract_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ValueError("服务端缺少 pypdf 依赖，无法解析 PDF") from exc

    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            logger.warning("PDF 单页抽取失败: %s", exc)
    return "\n".join(parts).strip()


def _extract_docx(raw: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ValueError("服务端缺少 python-docx 依赖，无法解析 DOCX") from exc

    document = docx.Document(io.BytesIO(raw))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs).strip()


def _extract_text(raw: bytes) -> str:
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return raw.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return ""


def extract_text_from_file(filename: str, raw: bytes) -> str:
    """按扩展名分发抽取逻辑。返回的文本会被截断到 MAX_TEXT_CHARS 字符。"""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text = _extract_pdf(raw)
    elif name.endswith(".docx"):
        text = _extract_docx(raw)
    else:
        text = _extract_text(raw)

    if not text:
        raise ValueError("未能从文件中抽取到有效文本，请确认文件格式")

    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n...[已截断，原文过长]..."
    return text


# ---------- LLM 抽取 ----------

def _build_field_schema() -> tuple[dict[str, Any], list[str], list[str]]:
    """把 CountryRegistry 里的国家信息序列化成 prompt 用的 schema。"""
    configs = CountryRegistry.get_all_configs()
    schema: dict[str, Any] = {}
    all_business_types: set[str] = set()
    country_codes: list[str] = []

    for code, cfg in configs.items():
        country_codes.append(code)
        all_business_types.update(cfg.business_types)
        schema[code] = {
            "country_name": cfg.country_name,
            "currency": cfg.currency,
            "business_fields": [
                {
                    "name": f.name,
                    "label": f.label,
                    "type": f.type,
                    "options": f.options,
                }
                for f in cfg.business_fields
            ],
        }
    return schema, sorted(all_business_types), country_codes


CONTRACT_EXTRACT_SYSTEM_PROMPT = """
你是一名税务合规审核助手，负责从跨境电商 / 出海业务的**合同文本**中抽取结构化字段，用于自动填充审核表单。

【任务】
- 只从原文中抽取真实存在的信息，禁止臆测。
- 输出必须是**严格 JSON**，不要包裹任何 markdown 代码块。
- 字段值若无法从原文明确判断，请输出 null / 空数组，不要猜。

【输出结构】
{
  "countries": ["TH", ...],              // 仅可选自 <allowed_country_codes>
  "business_type": "跨境电商零售",         // 仅可选自 <allowed_business_types>；不确定填 null
  "fields_by_country": {
    "TH": {
      "annual_sales": 12000000,
      "platforms": ["Shopee","Lazada"],
      "warehousing_mode": "海外仓",
      ...
    }
  },
  "confidence": { "TH.annual_sales": 0.9, "business_type": 0.8 }, // 0~1，可选
  "raw_hits": { "TH.annual_sales": "原文摘句" }                     // 抽取该字段所依据的原文，可选
}

【字段取值规则】
- 只输出 <country_schema> 中列出的字段 name，忽略其他。
- type=number：金额去除货币符号 / 千分位逗号后输出整数或浮点。中文单位（万 / 亿）需换算成本币数值。
- type=select：值必须**精确匹配** options 之一，匹配不到就置 null。
- type=multiselect：数组元素必须是 options 子集，未出现的选项不要塞进去。
- 若合同覆盖多个国家，为每个命中国家单独填一份 fields_by_country[code]。
""".strip()


def _coerce_number(value: Any) -> float | int | None:
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, str):
        s = value.replace(",", "").replace(" ", "")
        m = re.search(r"-?\d+(?:\.\d+)?", s)
        if m:
            num = float(m.group(0))
            return int(num) if num.is_integer() else num
    return None


def _sanitize_fields(
    fields_by_country: dict[str, Any],
) -> dict[str, dict[str, Any]]:
    """按 country_schema 白名单过滤 LLM 输出。"""
    configs = CountryRegistry.get_all_configs()
    cleaned: dict[str, dict[str, Any]] = {}

    if not isinstance(fields_by_country, dict):
        return cleaned

    for code, values in fields_by_country.items():
        code = str(code).upper()
        if code not in configs or not isinstance(values, dict):
            continue
        cfg = configs[code]
        cleaned_country: dict[str, Any] = {}
        for f in cfg.business_fields:
            if f.name not in values:
                continue
            raw = values[f.name]
            if raw is None or raw == "":
                continue

            if f.type == "number":
                num = _coerce_number(raw)
                if num is None:
                    continue
                if f.min_value is not None and num < f.min_value:
                    continue
                if f.max_value is not None and num > f.max_value:
                    continue
                cleaned_country[f.name] = num

            elif f.type == "select":
                if isinstance(raw, str) and raw in f.options:
                    cleaned_country[f.name] = raw

            elif f.type == "multiselect":
                if isinstance(raw, list):
                    picked = [item for item in raw if item in f.options]
                    if picked:
                        cleaned_country[f.name] = picked

            else:  # text or unknown
                if isinstance(raw, (str, int, float)):
                    cleaned_country[f.name] = raw

        if cleaned_country:
            cleaned[code] = cleaned_country
    return cleaned


async def extract_audit_fields(text: str) -> dict[str, Any]:
    """调用 LLM 抽取合同字段，返回已按注册表白名单校验的结构化数据。"""
    country_schema, business_types, country_codes = _build_field_schema()

    user_prompt = json.dumps(
        {
            "allowed_country_codes": country_codes,
            "allowed_business_types": business_types,
            "country_schema": country_schema,
            "contract_text": text,
        },
        ensure_ascii=False,
    )

    _model, content = await _achat_with_fallback(CONTRACT_EXTRACT_SYSTEM_PROMPT, user_prompt)

    try:
        parsed = _extract_json(content)
    except json.JSONDecodeError as exc:
        logger.warning("合同抽取 LLM 输出无法解析: %s", content[:500])
        raise ValueError("AI 抽取结果无法解析，请重试或人工填写") from exc

    # 顶层字段校验
    countries = parsed.get("countries") or []
    countries = [c for c in countries if isinstance(c, str) and c.upper() in country_codes]
    countries = list(dict.fromkeys(c.upper() for c in countries))  # 去重保序

    business_type = parsed.get("business_type")
    if business_type not in business_types:
        business_type = None

    fields_by_country = _sanitize_fields(parsed.get("fields_by_country") or {})

    # countries 未直接给出时，从 fields_by_country 的 key 推断
    if not countries and fields_by_country:
        countries = list(fields_by_country.keys())

    return {
        "countries": countries,
        "business_type": business_type,
        "fields_by_country": fields_by_country,
        "confidence": parsed.get("confidence") if isinstance(parsed.get("confidence"), dict) else {},
        "raw_hits": parsed.get("raw_hits") if isinstance(parsed.get("raw_hits"), dict) else {},
    }
